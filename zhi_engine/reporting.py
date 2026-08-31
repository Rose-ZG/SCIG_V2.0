from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _set_font(run, font_name: str = "Microsoft YaHei", size: int = 10, bold: bool = False, color: str | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size, bold, color in [
        ("Title", 20, True, "183153"),
        ("Heading 1", 14, True, "1F3D5B"),
        ("Heading 2", 12, True, "234B6A"),
        ("Heading 3", 11, True, "345D82"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)


def _add_bullet(doc: Document, text: str, level: int = 0) -> None:
    para = doc.add_paragraph(style="List Bullet")
    if level:
        para.paragraph_format.left_indent = Inches(0.28 * level)
    run = para.add_run(text)
    _set_font(run, size=10)


def _add_numbered(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    _set_font(run, size=10)


def _add_key_value_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in items:
        row = table.add_row().cells
        _set_cell_text(row[0], key, bold=True, color="1F3D5B", size=10)
        _set_cell_text(row[1], value, size=10)
        _set_cell_shading(row[0], "EAF2FB")
    doc.add_paragraph("")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        _set_cell_text(hdr[index], header, bold=True, color="FFFFFF", size=9)
        _set_cell_shading(hdr[index], "2F75B5")
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            _set_cell_text(cells[index], value, size=9)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths, strict=False):
                cell.width = Inches(width)


def build_report_docx(result: dict, title: str, logo_path: str | Path | None = None) -> bytes:
    doc = Document()
    _style_document(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    if logo_path:
        logo = Path(logo_path)
        if logo.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(logo), width=Inches(0.75))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_font(run, size=20, bold=True, color="163A5F")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"生成时间：{result.get('generated_at', '')}")
    _set_font(run, size=9, color="5D6D7E")

    doc.add_paragraph("")

    doc.add_heading("一、分析摘要", level=1)
    summary = result.get("summary", {})
    best = result.get("recommended_model", {})
    _add_key_value_table(
        doc,
        [
            ("样本数", str(summary.get("sample_count", "-"))),
            ("温度范围", f"{summary.get('temperature_range', ['-', '-'])[0]} - {summary.get('temperature_range', ['-', '-'])[1]} °C"),
            ("推荐模型", str(best.get("name", "-"))),
            ("拟合 R²", str(best.get("r2", "-"))),
            ("RMSE", str(best.get("rmse", "-"))),
            ("异常点", f"{summary.get('anomaly_count', 0)} 个"),
        ]
    )

    doc.add_heading("二、三大核心技术", level=1)
    for item in result.get("core_technologies", []):
        _add_bullet(doc, f"{item['title']}：{item['detail']}")

    physics = result.get("physics_constraints", {})
    physics_summary = physics.get("data_profile", {}).get("summary", {})
    reverse_solution = physics.get("reverse_solution") or []
    if isinstance(reverse_solution, list) and reverse_solution:
        reverse_target = max(reverse_solution, key=lambda item: abs(float(item.get("delta", 0))))
    elif isinstance(reverse_solution, dict):
        reverse_target = reverse_solution
    else:
        reverse_target = {}
    doc.add_heading("三、物理约束校验层", level=1)
    _add_key_value_table(
        doc,
        [
            ("物理评分", str(physics_summary.get("score", summary.get("physics_score", "-")))),
            ("边界越界", f"{physics_summary.get('boundary_violations', 0)} 个"),
            ("单调违例", f"{physics_summary.get('trend_violation_count', 0)} 个"),
            ("平均投影修正", str(physics_summary.get("projection_gap_mean", 0))),
            ("最大投影修正", str(physics_summary.get("projection_gap_max", 0))),
            ("最大修正温度", f"{reverse_target.get('temperature', '-')} °C"),
            ("最大修正量", str(reverse_target.get("delta", 0))),
        ],
    )
    for rule in physics.get("rules", []):
        _add_bullet(doc, f"{rule.get('title', '')}：{rule.get('detail', '')}")

    doc.add_heading("四、LLM + 符号回归候选", level=1)
    models = result.get("models", [])
    _add_table(
        doc,
        ["模型", "公式", "R²", "RMSE", "约束评分"],
        [
            [
                m.get("name", ""),
                m.get("equation", ""),
                str(m.get("r2", "")),
                str(m.get("rmse", "")),
                str(m.get("constraint", {}).get("score", "")),
            ]
            for m in models
        ],
        widths=[1.15, 3.05, 0.58, 0.68, 0.82],
    )

    doc.add_heading("五、假设性排行层", level=1)
    ranking_items = result.get("hypothesis_ranking", {}).get("items") or physics.get("rankings", [])
    if ranking_items:
        _add_table(
            doc,
            ["排名", "假设模型", "综合评分", "约束评分", "是否可行"],
            [
                [
                    str(item.get("rank", "")),
                    str(item.get("name", "")),
                    str(item.get("combined_score", "")),
                    str(item.get("constraint_score", "")),
                    "是" if item.get("feasible") else "否",
                ]
                for item in ranking_items
            ],
            widths=[0.55, 1.6, 1.0, 1.0, 0.8],
        )
    else:
        _add_bullet(doc, "暂无可展示的假设排行。")

    doc.add_heading("六、推荐模型", level=1)
    _add_bullet(doc, f"方程：{best.get('equation', '-')}")
    params = best.get("parameters", {})
    if params:
        _add_bullet(doc, "参数：" + ", ".join(f"{k}={v}" for k, v in params.items()))
    _add_bullet(doc, f"物理说明：{best.get('physics', '-')}")

    doc.add_heading("七、异常点", level=1)
    anomalies = result.get("anomalies", [])
    if anomalies:
        _add_table(
            doc,
            ["温度", "转化率", "批次", "等级", "说明"],
            [[str(a["temperature"]), str(a["conversion"]), str(a["batch"]), str(a["severity"]), a["reason"]] for a in anomalies],
            widths=[0.8, 0.8, 0.8, 0.6, 3.8],
        )
    else:
        _add_bullet(doc, "当前未发现需要重点复核的异常点。")

    doc.add_heading("八、下一步实验建议", level=1)
    for item in result.get("suggestions", []):
        _add_numbered(doc, item)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
